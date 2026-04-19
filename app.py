import streamlit as st
import google.generativeai as genai
import time
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
import re

# --- Konfigurasi Halaman ---
st.set_page_config(
    page_title="AI Lesson Plan Generator",
    page_icon="📚",
    layout="centered"
)

# --- Header ---
st.title("📚 AI Lesson Plan Generator")
st.caption("Powered by Google Gemini AI | Generate RPP berkualitas dengan muatan kearifan lokal Jawa Timur")

# --- Ambil API Key dari Streamlit Secrets ---
try:
    gemini_api_key = st.secrets["GEMINI_API_KEY"]
except:
    st.error("❌ API Key tidak ditemukan. Setup Secrets terlebih dahulu!")
    st.stop()

# --- Konfigurasi Gemini ---
MODEL_NAME = "gemini-2.5-flash"
genai.configure(api_key=gemini_api_key)

# --- Session State ---
if 'generated_content' not in st.session_state:
    st.session_state.generated_content = None
if 'is_generated' not in st.session_state:
    st.session_state.is_generated = False
if 'current_topic' not in st.session_state:
    st.session_state.current_topic = ""

# --- Dropdown Kelas ---
kelas_options = [
    "Kelas 1 SD", "Kelas 2 SD", "Kelas 3 SD", "Kelas 4 SD", "Kelas 5 SD", "Kelas 6 SD",
    "Kelas VII SMP", "Kelas VIII SMP", "Kelas IX SMP",
    "Kelas X SMA", "Kelas XI SMA", "Kelas XII SMA"
]


# ══════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════

def _add_inline_runs(paragraph, text):
    """Parse ***bold-italic***, **bold**, *italic* and add proper Word runs."""
    pattern = r'(\*{3}.+?\*{3}|\*{2}.+?\*{2}|\*.+?\*)'
    parts = re.split(pattern, text, flags=re.DOTALL)
    for part in parts:
        if not part:
            continue
        if part.startswith('***') and part.endswith('***'):
            r = paragraph.add_run(part[3:-3])
            r.bold = True
            r.italic = True
        elif part.startswith('**') and part.endswith('**'):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            r = paragraph.add_run(part[1:-1])
            r.italic = True
        else:
            paragraph.add_run(part)


def _strip_markdown(text):
    """Remove bold/italic markers for headings and table headers."""
    text = re.sub(r'\*{1,3}(.*?)\*{1,3}', r'\1', text)
    text = re.sub(r'_{1,3}(.*?)_{1,3}', r'\1', text)
    return text.strip()


def _is_table_row(line):
    """True if line looks like  | cell | cell |"""
    s = line.strip()
    return s.startswith('|') and s.endswith('|') and len(s) > 2


def _is_separator_row(line):
    """True for alignment rows like  |---|:---:|:---|"""
    s = line.strip()
    if not _is_table_row(s):
        return False
    inner = s[1:-1]
    return all(re.match(r'^[\s\-:]+$', c) for c in inner.split('|'))


def _parse_table_cells(line):
    """Return list of cell text strings from one markdown table row."""
    s = line.strip().strip('|')
    return [c.strip() for c in s.split('|')]


def _set_cell_border(cell):
    """Apply thin grey border to every edge of a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    '4')        # ½ pt
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'AAAAAA')
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _shade_cell(cell, fill_hex):
    """Set cell background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    tcPr.append(shd)


def _add_word_table(doc, header_cells, data_rows):
    """
    Render a proper Word table from parsed markdown table data.
    header_cells : list[str]        — column headers
    data_rows    : list[list[str]]  — body rows
    """
    n_cols = max(len(header_cells), max((len(r) for r in data_rows), default=1))
    n_rows = 1 + len(data_rows)

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = 'Table Grid'

    # Header row — blue background, bold text
    for col_i, cell_text in enumerate(header_cells):
        cell = table.rows[0].cells[col_i]
        _set_cell_border(cell)
        _shade_cell(cell, 'D9E1F2')
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(_strip_markdown(cell_text))
        run.bold = True
        run.font.size = Pt(10)

    # Data rows — alternating white / light grey (zebra)
    for row_i, row_cells in enumerate(data_rows):
        fill = 'F2F2F2' if row_i % 2 == 0 else 'FFFFFF'
        for col_i in range(n_cols):
            cell_text = row_cells[col_i] if col_i < len(row_cells) else ''
            cell = table.rows[row_i + 1].cells[col_i]
            _set_cell_border(cell)
            _shade_cell(cell, fill)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            _add_inline_runs(p, cell_text)
            for run in p.runs:
                run.font.size = Pt(10)

    doc.add_paragraph()   # space after table


# ══════════════════════════════════════════════════════════════
# MAIN: create_word_document
# ══════════════════════════════════════════════════════════════

def create_word_document(content, topic_name):
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    # Title
    title_para = doc.add_heading('RENCANA PELAKSANAAN PEMBELAJARAN (RPP)', level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f"Topik: {topic_name}")
    r.bold = True

    doc.add_paragraph('─' * 60)

    # ── Group lines into ('line', text) or ('table', [lines]) segments ──
    lines    = content.split('\n')
    segments = []
    i = 0
    while i < len(lines):
        if _is_table_row(lines[i]):
            block = []
            while i < len(lines) and _is_table_row(lines[i]):
                block.append(lines[i])
                i += 1
            segments.append(('table', block))
        else:
            segments.append(('line', lines[i]))
            i += 1

    # ── Render each segment ─────────────────────────────────────────────
    for seg_type, seg_data in segments:

        # ── TABLE ──────────────────────────────────────────────────────
        if seg_type == 'table':
            header_cells = None
            data_rows    = []
            for tl in seg_data:
                if _is_separator_row(tl):
                    continue                    # skip |---|---| separator
                cells = _parse_table_cells(tl)
                if header_cells is None:
                    header_cells = cells        # first content row = header
                else:
                    data_rows.append(cells)
            if header_cells:
                _add_word_table(doc, header_cells, data_rows)
            continue

        # ── NORMAL LINE ─────────────────────────────────────────────────
        raw      = seg_data
        stripped = raw.strip()

        # Empty line
        if not stripped:
            doc.add_paragraph()
            continue

        # Heading  # / ## / ### / ####
        m = re.match(r'^(#{1,4})\s+(.*)', stripped)
        if m:
            doc.add_heading(_strip_markdown(m.group(2)), level=min(len(m.group(1)), 4))
            continue

        # Numbered list  1. or 1)
        m = re.match(r'^(\d+)[.)]\s+(.*)', stripped)
        if m:
            para = doc.add_paragraph(style='List Number')
            _add_inline_runs(para, m.group(2))
            continue

        # Indented bullet (2+ leading spaces before - * +)
        m = re.match(r'^\s{2,}[-*+]\s+(.*)', raw)
        if m:
            para = doc.add_paragraph(style='List Bullet 2')
            _add_inline_runs(para, m.group(1))
            continue

        # Bullet  - / * / +
        m = re.match(r'^[-*+]\s+(.*)', stripped)
        if m:
            indent = len(raw) - len(raw.lstrip())
            style  = 'List Bullet 2' if indent >= 4 else 'List Bullet'
            para   = doc.add_paragraph(style=style)
            _add_inline_runs(para, m.group(1))
            continue

        # Horizontal rule  --- / *** / ___
        if re.match(r'^[-*_]{3,}$', stripped):
            doc.add_paragraph('─' * 60)
            continue

        # Standalone bold line  **text**  → sub-heading style
        if stripped.startswith('**') and stripped.endswith('**') \
                and stripped.count('**') == 2 and len(stripped) > 4:
            para = doc.add_paragraph()
            r    = para.add_run(_strip_markdown(stripped))
            r.bold = True
            r.font.size = Pt(12)
            continue

        # Plain paragraph (with inline bold/italic)
        para = doc.add_paragraph()
        _add_inline_runs(para, stripped)

    # Footer
    from datetime import datetime
    footer      = doc.sections[0].footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.clear()
    footer_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.add_run(
        f"Generated by AI Lesson Plan Generator | {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    )

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ══════════════════════════════════════════════════════════════
# Gemini caller — with automatic continuation if output is cut off
# ══════════════════════════════════════════════════════════════

def call_gemini(prompt_text):
    model        = genai.GenerativeModel(MODEL_NAME)
    gen_config   = {"temperature": 0.4, "max_output_tokens": 8000}
    MAX_CHUNKS   = 5          # safety cap — at most 5 continuation calls
    full_content = ""

    # Build a mutable conversation history for multi-turn continuation
    history = [{"role": "user", "parts": [prompt_text]}]

    for chunk_num in range(MAX_CHUNKS):
        # Retry loop for transient errors (503, high demand)
        response = None
        for attempt in range(3):
            try:
                response = model.generate_content(
                    history,
                    generation_config=gen_config
                )
                break
            except Exception as e:
                if attempt < 2 and ("503" in str(e) or "high demand" in str(e).lower()):
                    time.sleep(2)
                    continue
                st.error(f"Error: {str(e)}")
                return None

        if response is None:
            return None

        # Extract text from this chunk
        chunk_text = response.text
        # Strip markdown code fences only on the very first chunk
        if chunk_num == 0:
            chunk_text = re.sub(r'^```\w*\n?', '', chunk_text)
        chunk_text = re.sub(r'\n?```$', '', chunk_text)

        full_content += chunk_text

        # Check finish reason — 'STOP' means model finished naturally
        finish_reason = None
        try:
            finish_reason = response.candidates[0].finish_reason.name
        except Exception:
            pass

        # If the model finished on its own, we're done
        if finish_reason != "MAX_TOKENS":
            break

        # Output was cut off — add assistant reply to history and ask to continue
        history.append({"role": "model",  "parts": [chunk_text]})
        history.append({"role": "user",   "parts": ["Lanjutkan tepat dari titik berhenti, tanpa pengulangan."]})

    return full_content.strip()


# ══════════════════════════════════════════════════════════════
# UI — Form
# ══════════════════════════════════════════════════════════════

with st.form("lesson_form"):
    topic = st.text_area(
        "📖 **Topik Pembelajaran**",
        height=80,
        placeholder="Contoh: Memperkenalkan Tari Remo kepada siswa",
        help="Minimal 3 kata, maksimal 100 kata (MAX 500 KARAKTER)",
        max_chars=500
    )

    col1, col2 = st.columns(2)
    with col1:
        language = st.selectbox("🌐 Bahasa", ["Indonesian", "English"], index=0)
        grade    = st.selectbox("🎓 Kelas", kelas_options, index=4)
    with col2:
        std1 = st.text_area(
            "🎯 Standar 1 (Capaian Pembelajaran)",
            height=60,
            help="Minimal 2 kata, maksimal 25 kata (MAX 125 KARAKTER)",
            max_chars=125
        )
        std2 = st.text_area(
            "📝 Standar 2 (Tujuan Pembelajaran)",
            height=60,
            help="Minimal 2 kata, maksimal 25 kata (MAX 125 KARAKTER)",
            max_chars=125
        )

    time_minutes = st.text_input(
        "⏰ Alokasi Waktu (Menit)",
        placeholder="Contoh: 70",
        help="Hanya angka, maksimal 3 digit (contoh: 70, 90, 120)",
        max_chars=3
    )

    submitted = st.form_submit_button("🚀 Generate Lesson Plan", use_container_width=True)


# ── Validation & generation ─────────────────────────────────
if submitted:
    valid = True

    if not topic.strip():
        st.error("❌ Topik tidak boleh kosong!"); valid = False
    else:
        wc = len(topic.strip().split())
        if   wc < 3:   st.error(f"❌ Topik minimal 3 kata! Saat ini: {wc} kata");   valid = False
        elif wc > 100: st.error(f"❌ Topik maksimal 100 kata! Saat ini: {wc} kata"); valid = False

    if not std1.strip():
        st.error("❌ Standar 1 tidak boleh kosong!"); valid = False
    else:
        wc = len(std1.strip().split())
        if   wc < 2:  st.error(f"❌ Standar 1 minimal 2 kata! Saat ini: {wc} kata");   valid = False
        elif wc > 25: st.error(f"❌ Standar 1 maksimal 25 kata! Saat ini: {wc} kata"); valid = False

    if not std2.strip():
        st.error("❌ Standar 2 tidak boleh kosong!"); valid = False
    else:
        wc = len(std2.strip().split())
        if   wc < 2:  st.error(f"❌ Standar 2 minimal 2 kata! Saat ini: {wc} kata");   valid = False
        elif wc > 25: st.error(f"❌ Standar 2 maksimal 25 kata! Saat ini: {wc} kata"); valid = False

    if not time_minutes.strip():
        st.error("❌ Alokasi waktu tidak boleh kosong!"); valid = False
    else:
        if not time_minutes.isdigit():
            st.error("❌ Alokasi waktu harus berupa ANGKA saja (contoh: 70)"); valid = False
        elif int(time_minutes) < 10:
            st.error("❌ Alokasi waktu minimal 10 menit!"); valid = False

    if valid:
        st.session_state.current_topic = topic
        with st.spinner("AI sedang menyusun Rencana Pembelajaran..."):
            prompt = f"""
            Buat Rencana Pelaksanaan Pembelajaran (RPP) dengan detail berikut:

            Topik: {topic}
            Bahasa: {language}
            Kelas: {grade}
            Standar 1: {std1}
            Standar 2: {std2}
            Alokasi Waktu: {time_minutes} Menit

            Integrasikan kearifan lokal Jawa Timur.

            Struktur:
            1. KOMPETENSI AWAL
            2. PROFIL PELAJAR PANCASILA
            3. KEGIATAN PEMBELAJARAN (Pendahuluan, Inti, Penutup)
            4. PENILAIAN

            Gunakan format:
            ## untuk judul utama
            ### untuk sub judul
            - untuk list
            **teks** untuk penekanan
            | tabel | jika diperlukan |

            Langsung ke konten, tanpa kata pengantar.
            """
            result = call_gemini(prompt)
            if result:
                st.session_state.generated_content = result
                st.session_state.is_generated      = True
                st.success("✅ Lesson Plan berhasil dibuat!")
            else:
                st.error("Gagal menghasilkan. Silakan coba lagi.")


# ── Output ───────────────────────────────────────────────────
if st.session_state.is_generated and st.session_state.generated_content:
    st.divider()

    doc_file = create_word_document(
        st.session_state.generated_content,
        st.session_state.current_topic[:50]
    )

    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.download_button(
            label="📥 Download Lesson Plan (Word)",
            data=doc_file,
            file_name=f"Lesson_Plan_{int(time.time())}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

    st.divider()
    st.subheader("📄 Preview")
    st.markdown(st.session_state.generated_content, unsafe_allow_html=True)
    st.caption("💡 Download file Word untuk hasil cetak yang rapi.")
